"""
Génère le guide d'installation de Kobo Calibre Browser en .docx
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT = "Guide_Installation_Kobo_Calibre_Browser.docx"

# ── Palette de couleurs ────────────────────────────────────────────────────────
BLUE_DARK   = RGBColor(0x1F, 0x3D, 0x6E)   # Titres principaux
BLUE_MED    = RGBColor(0x2E, 0x6D, 0xAE)   # Titres secondaires
BLUE_LIGHT  = RGBColor(0xDE, 0xEB, 0xF7)   # Fond entête de tableau
GREY_LIGHT  = RGBColor(0xF5, 0xF5, 0xF5)   # Fond code / note
GREY_BORDER = RGBColor(0xBB, 0xBB, 0xBB)   # Bordures
GREEN       = RGBColor(0x1A, 0x7A, 0x3C)   # Succès / validation
ORANGE      = RGBColor(0xC5, 0x55, 0x00)   # Avertissement
BLACK       = RGBColor(0x00, 0x00, 0x00)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)

# ── Helpers XML ───────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="BBBBBB"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def set_row_height(row, height_cm: float):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(height_cm * 567)))  # 567 twips/cm
    trPr.append(trHeight)

def add_bookmark(para, bookmark_id: int, bookmark_name: str):
    """Ajoute un signet sur un paragraphe (pour le sommaire)."""
    run = para.runs[0] if para.runs else para.add_run()
    bStart = OxmlElement("w:bookmarkStart")
    bStart.set(qn("w:id"), str(bookmark_id))
    bStart.set(qn("w:name"), bookmark_name)
    bEnd = OxmlElement("w:bookmarkEnd")
    bEnd.set(qn("w:id"), str(bookmark_id))
    run._r.addprevious(bStart)
    run._r.addnext(bEnd)

# ── Constructeurs de paragraphes ──────────────────────────────────────────────

def heading1(doc, text, bookmark_id=None, bookmark_name=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    # Trait supérieur
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "12")
    top.set(qn("w:space"), "4")
    top.set(qn("w:color"), "1F3D6E")
    pBdr.append(top)
    pPr.append(pBdr)
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = BLUE_DARK
    run.font.name = "Calibri"
    if bookmark_id is not None:
        add_bookmark(p, bookmark_id, bookmark_name)
    return p

def heading2(doc, text, bookmark_id=None, bookmark_name=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = BLUE_MED
    run.font.name = "Calibri"
    if bookmark_id is not None:
        add_bookmark(p, bookmark_id, bookmark_name)
    return p

def heading3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK
    run.font.name = "Calibri"
    return p

def body(doc, text, space_after=Pt(6)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = BLACK
    run.font.name = "Calibri"
    return p

def body_bold(doc, text, suffix=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(text)
    r1.font.bold = True
    r1.font.size = Pt(10.5)
    r1.font.name = "Calibri"
    if suffix:
        r2 = p.add_run(suffix)
        r2.font.size = Pt(10.5)
        r2.font.name = "Calibri"
    return p

def bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    indent = Cm(0.5 + level * 0.7)
    p.paragraph_format.left_indent = indent
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.font.bold = True
        r1.font.size = Pt(10.5)
        r1.font.name = "Calibri"
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    r2.font.name = "Calibri"
    return p

def numbered(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.font.bold = True
        r1.font.size = Pt(10.5)
        r1.font.name = "Calibri"
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    r2.font.name = "Calibri"
    return p

def code_block(doc, lines: list[str], caption=None):
    """Bloc de code avec fond gris."""
    if caption:
        pc = doc.add_paragraph()
        pc.paragraph_format.space_after = Pt(2)
        rc = pc.add_run(caption)
        rc.font.italic = True
        rc.font.size = Pt(9)
        rc.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        rc.font.name = "Calibri"

    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.right_indent = Cm(0.5)
        # Fond gris
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F0F0F0")
        pPr.append(shd)
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    # Espace après le bloc
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(6)

def note_box(doc, text, kind="info"):
    """Encadré info/avertissement."""
    colors = {
        "info":    ("DEF0FA", "0070C0", "ℹ️  Note"),
        "warning": ("FFF3CD", "C55500", "⚠️  Attention"),
        "success": ("D6F0E0", "1A7A3C", "✅  Conseil"),
    }
    fill, border_hex, label = colors.get(kind, colors["info"])
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, fill)
    set_cell_borders(cell, border_hex)
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_before = Pt(4)
    p1.paragraph_format.space_after = Pt(2)
    r_label = p1.add_run(label + "  ")
    r_label.font.bold = True
    r_label.font.size = Pt(10)
    r_label.font.name = "Calibri"
    r_body = p1.add_run(text)
    r_body.font.size = Pt(10)
    r_body.font.name = "Calibri"
    cell.add_paragraph().paragraph_format.space_after = Pt(4)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)

def two_col_table(doc, headers, rows, col_widths_cm=None):
    """Tableau à N colonnes."""
    n = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # En-tête
    header_row = tbl.rows[0]
    set_row_height(header_row, 0.8)
    for i, h in enumerate(headers):
        cell = header_row.cells[i]
        set_cell_bg(cell, "1F3D6E")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(h)
        r.font.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(10)
        r.font.name = "Calibri"

    # Lignes
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        bg = "F5F5F5" if ri % 2 == 0 else "FFFFFF"
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            set_cell_borders(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            # Texte bold+normal si séparé par |
            if "|" in str(cell_text):
                parts = cell_text.split("|", 1)
                rb = p.add_run(parts[0])
                rb.font.bold = True
                rb.font.size = Pt(9.5)
                rb.font.name = "Courier New"
                rn = p.add_run(parts[1])
                rn.font.size = Pt(9.5)
                rn.font.name = "Calibri"
            else:
                r = p.add_run(str(cell_text))
                r.font.size = Pt(9.5)
                r.font.name = "Calibri"

    # Largeurs de colonnes
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)
    return tbl

def page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)

def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)

# ══════════════════════════════════════════════════════════════════════════════
# CONSTRUCTION DU DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

doc = Document()

# Marges
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── PAGE DE COUVERTURE ────────────────────────────────────────────────────────
p_title = doc.add_paragraph()
p_title.paragraph_format.space_before = Cm(3)
p_title.paragraph_format.space_after = Pt(8)
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_title.add_run("📚  Kobo Calibre Browser")
r.font.bold = True
r.font.size = Pt(28)
r.font.color.rgb = BLUE_DARK
r.font.name = "Calibri"

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_after = Pt(6)
r2 = p_sub.add_run("Guide d'installation et de configuration")
r2.font.size = Pt(16)
r2.font.color.rgb = BLUE_MED
r2.font.name = "Calibri"

p_ver = doc.add_paragraph()
p_ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_ver.paragraph_format.space_after = Cm(2)
r3 = p_ver.add_run("Version 1.0  ·  Serveur Unraid  ·  Traefik + Authelia")
r3.font.size = Pt(11)
r3.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
r3.font.name = "Calibri"
r3.font.italic = True

# Tableau résumé stack
two_col_table(doc,
    ["Composant", "Rôle", "Accès"],
    [
        ["Unraid + Docker",       "Serveur hôte, orchestration des containers",      "Local / SSH"],
        ["Calibre",               "Gestionnaire de bibliothèque ebook",               "GUI desktop ou Web"],
        ["kobo-calibre-browser",  "Application web de navigation (ce projet)",        "https://kobo.majurel.fr"],
        ["calibre-web",           "Interface Calibre existante (Kobo Sync actif)",    "https://calibre.majurel.fr"],
        ["Traefik",               "Reverse-proxy HTTPS, routage des sous-domaines",   "Tableau de bord Traefik"],
        ["Authelia",              "Authentification SSO (2FA, One-Time PIN)",          "https://auth.majurel.fr"],
        ["Kobo Clara 2E",         "Liseuse e-ink — navigateur intégré",               "WiFi / 4G"],
    ],
    col_widths_cm=[4.5, 8.0, 4.5]
)

page_break(doc)

# ── SOMMAIRE ──────────────────────────────────────────────────────────────────
h = doc.add_paragraph()
h.paragraph_format.space_after = Pt(16)
r = h.add_run("Sommaire")
r.font.bold = True
r.font.size = Pt(20)
r.font.color.rgb = BLUE_DARK
r.font.name = "Calibri"

toc_entries = [
    ("1.", "Prérequis et architecture",        1),
    ("2.", "Préparer la bibliothèque Calibre", 1),
    ("3.", "Déployer kobo-calibre-browser",    1),
    ("  3.1", "Cloner le dépôt",               2),
    ("  3.2", "Configurer l'environnement",    2),
    ("  3.3", "Lancer le container",           2),
    ("4.", "Configurer Traefik",               1),
    ("5.", "Configurer Authelia",              1),
    ("6.", "Configurer le DNS",                1),
    ("7.", "Configurer la Kobo Clara 2E",      1),
    ("8.", "Vérification et tests",            1),
    ("9.", "Dépannage",                        1),
]

for num, title, level in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm((level - 1) * 0.8)
    r_num = p.add_run(num + "  ")
    r_num.font.bold = level == 1
    r_num.font.size = Pt(10.5 if level == 1 else 10)
    r_num.font.color.rgb = BLUE_MED if level == 1 else RGBColor(0x44, 0x44, 0x44)
    r_num.font.name = "Calibri"
    r_title = p.add_run(title)
    r_title.font.bold = level == 1
    r_title.font.size = Pt(10.5 if level == 1 else 10)
    r_title.font.color.rgb = BLACK
    r_title.font.name = "Calibri"

page_break(doc)

# ── 1. PRÉREQUIS ET ARCHITECTURE ─────────────────────────────────────────────
heading1(doc, "1. Prérequis et architecture", 1, "prereqs")

body(doc, "Cette solution déploie une application web FastAPI sur un serveur Unraid. "
     "La bibliothèque Calibre existante est montée en lecture seule dans le container. "
     "Traefik expose l'application en HTTPS et Authelia contrôle l'authentification avant "
     "que toute requête n'atteigne le serveur.")

heading2(doc, "Prérequis techniques")

two_col_table(doc,
    ["Composant", "Version minimale", "Remarque"],
    [
        ["Unraid",   "6.12+",  "Community Apps activé recommandé"],
        ["Docker",   "24+",    "Inclus dans Unraid"],
        ["Traefik",  "v2 ou v3", "Réseau externe 'docker_network' configuré, entrypoint 'websecure' actif"],
        ["Authelia", "4.38+",  "Middleware 'authelia@file' déclaré (file-provider)"],
        ["Git",      "2.x",    "Pour cloner le dépôt"],
        ["Calibre",  "7.x",    "Base de données metadata.db existante avec livres importés"],
    ],
    col_widths_cm=[4.5, 3.5, 9.0]
)

heading2(doc, "Architecture réseau")

body(doc, "Le schéma suivant décrit le flux réseau entre la Kobo et la bibliothèque :")

code_block(doc, [
    "Kobo Clara 2E (navigateur)                  Serveur Unraid (réseau local / Internet)",
    "       │                                            │",
    "       │  HTTPS  https://kobo.majurel.fr           │",
    "       ├──────────────────────────────────────────►│  [Traefik :443]",
    "       │                                            │       │ vérifie middleware",
    "       │                                            │  [Authelia]",
    "       │  ← cookie de session si auth OK ──────────┤       │",
    "       │                                            │  [kobo-calibre :8000]",
    "       │                                            │       │ lit en lecture seule",
    "       │                                            │  [Volume Calibre Library]",
    "       │  HTML / covers / téléchargement EPUB ◄─────┤",
])

note_box(doc,
    "cloudflared (Cloudflare Tunnel) n'est PAS nécessaire ici : Traefik + "
    "Let's Encrypt gèrent déjà le HTTPS et l'exposition publique via le DNS.",
    "success")

page_break(doc)

# ── 2. BIBLIOTHÈQUE CALIBRE ──────────────────────────────────────────────────
heading1(doc, "2. Préparer la bibliothèque Calibre", 2, "calibre")

body(doc, "calibre-web est déjà déployé sur https://calibre.majurel.fr avec le Kobo Sync actif. "
     "kobo-calibre-browser lit la même bibliothèque en lecture seule — aucune action "
     "n'est nécessaire sur calibre-web, les deux services coexistent sans conflit.")

heading2(doc, "Localiser le chemin de la bibliothèque")
body(doc, "Vérifier le chemin dans le container Calibre existant :")
code_block(doc, [
    "# Sur l'hôte Unraid (shell ou terminal Unraid)",
    "docker inspect calibre | grep -i 'source\\|destination'",
    "",
    "# Chemin sur ce serveur :",
    "/mnt/user/data/media/ebooks/calibre_library",
], "Commande shell Unraid")

heading2(doc, "Structure attendue de la bibliothèque")
body(doc, "Le container kobo-calibre-browser attend la structure standard Calibre :")
code_block(doc, [
    "/mnt/user/data/media/ebooks/calibre_library/",
    "├── metadata.db                   ← Base SQLite (sera montée en read-only)",
    "├── Auteur, Prénom/",
    "│   └── Titre du livre (123)/",
    "│       ├── cover.jpg             ← Couverture",
    "│       ├── Titre du livre.epub   ← Fichier ebook",
    "│       └── Titre du livre.mobi   ← (optionnel)",
    "└── ...                           ← Autres auteurs/livres",
], "Structure Calibre standard")

note_box(doc,
    "Le container monte la bibliothèque en lecture seule (read_only: true). "
    "Calibre Desktop / calibre-web continuent de fonctionner normalement sans conflit.",
    "success")

page_break(doc)

# ── 3. DÉPLOYER KOBO-CALIBRE-BROWSER ─────────────────────────────────────────
heading1(doc, "3. Déployer kobo-calibre-browser", 3, "deploy")

heading2(doc, "3.1  Cloner le dépôt", 4, "clone")
body(doc, "Sur le serveur Unraid, ouvrir un terminal (via SSH ou l'interface web) :")
code_block(doc, [
    "# Se connecter en SSH à Unraid",
    "ssh root@<IP_UNRAID>",
    "",
    "# Créer le dossier d'installation",
    "mkdir -p /mnt/user/appdata/kobo-calibre-browser",
    "cd /mnt/user/appdata/kobo-calibre-browser",
    "",
    "# Cloner le dépôt",
    "git clone https://github.com/<votre-repo>/kobo-calibre-browser .",
], "Terminal SSH Unraid")

heading2(doc, "3.2  Configurer l'environnement", 5, "envconfig")
body(doc, "Copier le fichier d'exemple et renseigner les valeurs :")
code_block(doc, [
    "cp .env.example .env",
    "vi .env    # ou nano .env",
], "Copier la configuration")

body(doc, "Contenu du fichier .env à adapter :")
code_block(doc, [
    "# Chemin absolu vers la bibliothèque Calibre sur l'hôte",
    "CALIBRE_LIBRARY_PATH=/mnt/user/data/media/ebooks/calibre_library",
    "",
    "# Sous-domaine public de l'application",
    "KOBO_DOMAIN=kobo.majurel.fr",
    "",
    "# Réseau Docker externe de Traefik (vérifier avec : docker network ls)",
    "TRAEFIK_NETWORK=docker_network",
    "",
    "# Résolveur de certificat TLS configuré dans Traefik",
    "CERT_RESOLVER=cloudflare",
    "",
    "# Middleware Authelia (authelia@file si file-provider, authelia@docker sinon)",
    "AUTHELIA_MIDDLEWARE=authelia@file",
    "",
    "# Livres par page (défaut 20)",
    "ITEMS_PER_PAGE=20",
], "Fichier .env")

two_col_table(doc,
    ["Variable", "Description", "Valeur par défaut"],
    [
        ["CALIBRE_LIBRARY_PATH", "Chemin absolu de la bibliothèque Calibre sur l'hôte Unraid",  "à remplir obligatoirement"],
        ["KOBO_DOMAIN",          "Sous-domaine public HTTPS de l'application",                   "kobo.majurel.fr"],
        ["TRAEFIK_NETWORK",      "Nom du réseau Docker externe de Traefik",                      "docker_network"],
        ["CERT_RESOLVER",        "Nom du résolveur TLS déclaré dans Traefik",                    "cloudflare"],
        ["AUTHELIA_MIDDLEWARE",  "Nom du middleware Authelia dans Traefik",                       "authelia@file"],
        ["ITEMS_PER_PAGE",       "Nombre de livres affichés par page",                           "20"],
        ["LOG_LEVEL",            "Niveau de log uvicorn (debug/info/warning/error)",              "info"],
    ],
    col_widths_cm=[4.5, 8.5, 4.0]
)

heading2(doc, "3.3  Lancer le container", 6, "launch")
code_block(doc, [
    "# Construction de l'image et démarrage",
    "docker compose up -d --build",
    "",
    "# Vérifier que le container est démarré",
    "docker compose ps",
    "",
    "# Consulter les logs",
    "docker compose logs -f kobo-calibre",
], "Démarrage Docker")

body(doc, "Le container effectue un healthcheck automatique sur /health toutes les 30 secondes. "
     "Attendre que le statut passe de 'starting' à 'healthy' (environ 15 secondes).")

code_block(doc, [
    "# Test direct depuis l'hôte Unraid",
    "curl http://localhost:8000/health",
    "",
    "# Réponse attendue :",
    '{  "status": "ok",  "database": "connected",  "library_path": "/calibre-library"  }',
], "Test de santé")

note_box(doc,
    "Si le statut database est 'error', vérifier que CALIBRE_LIBRARY_PATH pointe "
    "bien vers le dossier contenant metadata.db, et non vers un sous-dossier.",
    "warning")

page_break(doc)

# ── 4. CONFIGURER TRAEFIK ─────────────────────────────────────────────────────
heading1(doc, "4. Configurer Traefik", 7, "traefik")

body(doc, "Le fichier docker-compose.yml inclut déjà les labels Traefik nécessaires. "
     "Il n'y a normalement rien à modifier dans la configuration Traefik elle-même — "
     "les labels sur le container suffisent.")

heading2(doc, "Labels Traefik dans docker-compose.yml")
body(doc, "Ces labels sont déjà présents dans le fichier fourni :")
code_block(doc, [
    "labels:",
    '  - "traefik.enable=true"',
    '  - "traefik.http.routers.kobo-calibre.rule=Host(`kobo.majurel.fr`)"',
    '  - "traefik.http.routers.kobo-calibre.entrypoints=websecure"',
    '  - "traefik.http.routers.kobo-calibre.tls=true"',
    '  - "traefik.http.routers.kobo-calibre.tls.certresolver=cloudflare"',
    '  - "traefik.http.services.kobo-calibre.loadbalancer.server.port=8000"',
    '  - "traefik.http.routers.kobo-calibre.middlewares=authelia@file"',
], "Labels Traefik dans docker-compose.yml")

heading2(doc, "Réseau Docker partagé")
body(doc, "Le container doit être sur le même réseau Docker externe que Traefik :")
code_block(doc, [
    "# Vérifier le nom du réseau Traefik",
    "docker network ls | grep docker_network",
    "",
    "# Si le réseau s'appelle autrement, modifier .env :",
    "TRAEFIK_NETWORK=nom_du_reseau_traefik",
    "",
    "# Vérifier que Traefik est bien sur ce réseau",
    "docker inspect traefik | grep -A5 'Networks'",
], "Vérification du réseau")

heading2(doc, "Vérification du routage Traefik")
body(doc, "Dans le tableau de bord Traefik, vérifier que le router 'kobo-calibre' apparaît "
     "avec le statut 'Enabled' et un certificat TLS valide :")
bullet(doc, "Accéder au tableau de bord Traefik (ex: https://traefik.majurel.fr)")
bullet(doc, "Aller dans HTTP → Routers")
bullet(doc, "Vérifier que 'kobo-calibre@docker' est listé et vert")
bullet(doc, "Vérifier que le certificat TLS Let's Encrypt est valide")

note_box(doc,
    "Ce serveur utilise le résolveur 'cloudflare'. Si ton Traefik utilise 'letsencrypt' "
    "ou un autre résolveur, modifier la variable CERT_RESOLVER dans le .env.",
    "info")

page_break(doc)

# ── 5. CONFIGURER AUTHELIA ────────────────────────────────────────────────────
heading1(doc, "5. Configurer Authelia", 8, "authelia")

body(doc, "Authelia protège l'application avant que les requêtes n'atteignent le serveur. "
     "Il faut ajouter une règle d'accès dans la configuration Authelia pour le nouveau sous-domaine.")

heading2(doc, "Ajouter la règle d'accès")
body(doc, "Dans le fichier de configuration Authelia (généralement "
     "/mnt/user/appdata/authelia/configuration.yml), "
     "ajouter une règle dans la section access_control :")
code_block(doc, [
    "access_control:",
    "  default_policy: deny",
    "  rules:",
    "    # ... règles existantes ...",
    "",
    "    # Règle pour Kobo Calibre Browser",
    "    - domain: kobo.majurel.fr",
    "      policy: one_factor   # ou two_factor pour plus de sécurité",
    "      # Optionnel : restreindre à certains utilisateurs",
    "      # subject:",
    "      #   - 'user:fabien'",
], "configuration.yml Authelia")

heading2(doc, "Durée de session recommandée pour la Kobo")
body(doc, "Le navigateur de la Kobo Clara 2E gère les cookies mais peut les perdre en cas "
     "de redémarrage. Il est recommandé de configurer une durée de session longue pour "
     "éviter des ré-authentifications fréquentes :")
code_block(doc, [
    "session:",
    "  expiration: 720h    # 30 jours",
    "  inactivity: 168h    # 7 jours d'inactivité",
    "  remember_me: 2160h  # 90 jours si 'Se souvenir de moi' coché",
], "Durée de session dans configuration.yml")

heading2(doc, "Méthode d'authentification recommandée")
body(doc, "Pour la Kobo, la méthode One-Time PIN (OTP par email) est recommandée car "
     "le navigateur Kobo supporte bien les formulaires HTML simples :")
two_col_table(doc,
    ["Méthode", "Compatibilité Kobo", "Configuration requise"],
    [
        ["One-Time PIN (email OTP)",   "✅ Excellente",  "Serveur SMTP configuré dans Authelia"],
        ["TOTP (Google Authenticator)","✅ Bonne",       "Application TOTP sur smartphone"],
        ["Mot de passe seul",          "✅ Excellente",  "Aucune (one_factor policy)"],
        ["WebAuthn / passkeys",        "❌ Non supporté", "Non disponible sur le navigateur Kobo"],
    ],
    col_widths_cm=[5.0, 3.5, 8.5]
)

heading2(doc, "Appliquer la configuration")
code_block(doc, [
    "# Redémarrer Authelia pour appliquer les changements",
    "docker restart authelia",
    "",
    "# Vérifier les logs Authelia",
    "docker logs authelia --tail 50",
], "Redémarrage Authelia")

note_box(doc,
    "Ce serveur utilise authelia@file (file-provider). Si Authelia est déclaré comme "
    "container Docker avec labels Traefik, utiliser authelia@docker à la place.",
    "info")

page_break(doc)

# ── 6. CONFIGURER LE DNS ──────────────────────────────────────────────────────
heading1(doc, "6. Configurer le DNS", 9, "dns")

body(doc, "Le sous-domaine kobo.majurel.fr doit pointer vers le serveur Unraid "
     "(ou vers le reverse proxy qui reçoit les connexions entrantes).")

heading2(doc, "Enregistrement DNS à créer")
two_col_table(doc,
    ["Type", "Nom", "Valeur", "TTL"],
    [
        ["A",     "kobo",  "<IP publique de la box / serveur>",  "3600"],
        ["CNAME", "kobo",  "majurel.fr  (si IP déjà sur le domaine racine)",  "3600"],
    ],
    col_widths_cm=[2.0, 3.0, 9.5, 2.5]
)

heading2(doc, "Avec Cloudflare DNS (si le domaine est géré par Cloudflare)")
body(doc, "Si majurel.fr est géré par Cloudflare, ajouter le sous-domaine dans le dashboard :")
numbered(doc, "Se connecter sur dash.cloudflare.com")
numbered(doc, "Sélectionner le domaine majurel.fr")
numbered(doc, "DNS → Add record → Type A → Name: kobo → IPv4: <IP publique>")
numbered(doc, "Proxy status : mettre en 'DNS only' (nuage gris) si Traefik gère le TLS, "
              "ou 'Proxied' (nuage orange) si Cloudflare doit aussi proxifier")

note_box(doc,
    "Si la box est derrière une IP dynamique, configurer un client DDNS (DuckDNS, "
    "Cloudflare DDNS via container) pour mettre à jour l'enregistrement automatiquement.",
    "warning")

heading2(doc, "Vérification DNS")
code_block(doc, [
    "# Depuis n'importe quelle machine (pas le serveur lui-même)",
    "nslookup kobo.majurel.fr",
    "",
    "# Doit retourner l'IP publique du serveur Unraid",
    "# Exemple :",
    "# Name:    kobo.majurel.fr",
    "# Address: 90.123.45.67",
    "",
    "# Tester la connectivité HTTPS",
    "curl -I https://kobo.majurel.fr/health",
], "Vérification DNS et HTTPS")

page_break(doc)

# ── 7. CONFIGURER LA KOBO ─────────────────────────────────────────────────────
heading1(doc, "7. Configurer la Kobo Clara 2E", 10, "kobo")

body(doc, "Aucune application tierce ni modification du firmware n'est nécessaire. "
     "La Kobo se connecte à l'application via son navigateur web intégré.")

heading2(doc, "Activer le navigateur web sur la Kobo")
body(doc, "Sur la Kobo Clara 2E, le navigateur est disponible nativement :")
numbered(doc, "Depuis l'écran d'accueil de la Kobo, appuyer sur l'icône Menu (☰) en haut à droite")
numbered(doc, "Aller dans Bêta → Navigateur web")
numbered(doc, "Si 'Navigateur web' n'apparaît pas : Paramètres → Compte → Fonctions bêta → activer 'Navigateur web'")

heading2(doc, "Se connecter à la bibliothèque")
numbered(doc, "S'assurer que la Kobo est connectée au WiFi")
numbered(doc, "Ouvrir le navigateur web de la Kobo")
numbered(doc, "Saisir l'URL dans la barre d'adresse :")
code_block(doc, ["https://kobo.majurel.fr"])
numbered(doc, "La page Authelia s'affiche — s'identifier avec son compte")
numbered(doc, "Après authentification, la bibliothèque s'affiche")

heading2(doc, "Créer un favori (accès rapide)")
numbered(doc, "Une fois sur https://kobo.majurel.fr, appuyer sur l'icône étoile dans la barre d'adresse")
numbered(doc, "Le signet est sauvegardé et accessible depuis la page d'accueil du navigateur")

heading2(doc, "Télécharger un livre sur la Kobo")
numbered(doc, "Naviguer jusqu'à la fiche d'un livre")
numbered(doc, "Appuyer sur le bouton 'EPUB' ou 'MOBI'")
numbered(doc, "Le fichier se télécharge automatiquement")
numbered(doc, "La Kobo propose d'ouvrir le fichier avec la bibliothèque → confirmer")
numbered(doc, "Le livre apparaît dans la bibliothèque de la Kobo")

note_box(doc,
    "Privilégier le format EPUB pour la Kobo Clara 2E (meilleur rendu, support natif). "
    "Le format MOBI est aussi supporté. Éviter le PDF pour les livres de texte (pas de "
    "reflow sur e-ink).",
    "success")

two_col_table(doc,
    ["Format", "Compatibilité Kobo", "Recommandation"],
    [
        ["EPUB",  "✅ Natif — reflow adaptatif", "⭐ Recommandé"],
        ["MOBI",  "✅ Supporté",                 "Acceptable"],
        ["AZW3",  "✅ Supporté",                 "Acceptable"],
        ["PDF",   "⚠️ Pas de reflow",            "À éviter pour romans"],
        ["TXT",   "✅ Supporté",                 "Pour textes simples"],
    ],
    col_widths_cm=[2.5, 6.0, 8.5]
)

page_break(doc)

# ── 8. VÉRIFICATION ET TESTS ──────────────────────────────────────────────────
heading1(doc, "8. Vérification et tests", 11, "tests")

heading2(doc, "Checklist de déploiement")

checks = [
    ("Container démarré", "docker compose ps  →  statut 'healthy'"),
    ("Base Calibre accessible", "curl http://localhost:8000/health  →  database: 'connected'"),
    ("DNS résolu", "nslookup kobo.majurel.fr  →  retourne l'IP correcte"),
    ("HTTPS valide", "curl -I https://kobo.majurel.fr  →  HTTP/2 200 ou 302"),
    ("Authelia actif", "Naviguer vers https://kobo.majurel.fr  →  page de login Authelia"),
    ("Après auth : page d'accueil", "La bibliothèque s'affiche avec le nombre de livres"),
    ("Navigation auteurs/genres", "Les listes sont correctement peuplées depuis Calibre"),
    ("Couvertures affichées", "Les couvertures apparaissent dans la grille de livres"),
    ("Téléchargement EPUB", "Le fichier se télécharge et s'ouvre dans la Kobo"),
    ("Test depuis Kobo réelle", "Naviguer vers https://kobo.majurel.fr depuis le navigateur Kobo"),
]

for label, detail in checks:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.3)
    r_check = p.add_run("☐  ")
    r_check.font.size = Pt(12)
    r_check.font.name = "Calibri"
    r_bold = p.add_run(label + " — ")
    r_bold.font.bold = True
    r_bold.font.size = Pt(10.5)
    r_bold.font.name = "Calibri"
    r_detail = p.add_run(detail)
    r_detail.font.size = Pt(10)
    r_detail.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    r_detail.font.name = "Courier New"

heading2(doc, "Test de résolution depuis un PC (simuler la Kobo)")
body(doc, "Pour vérifier le rendu e-ink sur un PC avant de tester sur la vraie Kobo :")
numbered(doc, "Ouvrir Chrome ou Firefox")
numbered(doc, "Ouvrir les Outils de développement (F12)")
numbered(doc, "Activer la vue responsive (Ctrl+Shift+M dans Chrome)")
numbered(doc, "Définir une résolution custom : 1072 × 1448 px (portrait)")
numbered(doc, "Naviguer vers https://kobo.majurel.fr")

page_break(doc)

# ── 9. DÉPANNAGE ──────────────────────────────────────────────────────────────
heading1(doc, "9. Dépannage", 12, "troubleshoot")

two_col_table(doc,
    ["Symptôme", "Cause probable", "Solution"],
    [
        [
            "Container ne démarre pas",
            "CALIBRE_LIBRARY_PATH incorrect ou dossier inexistant",
            "Vérifier le chemin et que metadata.db existe dans le dossier"
        ],
        [
            "database: 'error' dans /health",
            "metadata.db absent ou permissions insuffisantes",
            "chmod +r sur metadata.db ; vérifier que le chemin inclut 'Calibre Library'"
        ],
        [
            "404 sur kobo.majurel.fr",
            "DNS non propagé ou Traefik ne voit pas le container",
            "Attendre la propagation DNS ; vérifier 'docker network inspect docker_network'"
        ],
        [
            "Certificat TLS invalide",
            "Let's Encrypt n'a pas pu délivrer le cert (port 80/443 non joignable)",
            "Vérifier que le port 443 est ouvert sur la box / le firewall"
        ],
        [
            "Authelia redirige en boucle",
            "Le cookie de session n'est pas sauvegardé par le navigateur Kobo",
            "Vérifier que le domaine Authelia et kobo.majurel.fr sont sur le même domaine parent"
        ],
        [
            "Couvertures absentes",
            "cover.jpg absent pour certains livres dans Calibre",
            "Dans Calibre Desktop : sélectionner les livres → Éditer métadonnées → Télécharger couverture"
        ],
        [
            "Téléchargement EPUB échoue (404)",
            "Format non présent en base ou fichier manquant sur disque",
            "Vérifier dans Calibre que le fichier existe ; re-convertir si nécessaire"
        ],
        [
            "Page trop petite / trop grande sur Kobo",
            "Zoom navigateur Kobo modifié",
            "Dans le navigateur Kobo : icône paramètres → Zoom → 100%"
        ],
    ],
    col_widths_cm=[4.5, 5.5, 7.0]
)

heading2(doc, "Commandes de diagnostic utiles")
code_block(doc, [
    "# Logs en temps réel",
    "docker compose logs -f kobo-calibre",
    "",
    "# Vérifier les variables d'environnement injectées",
    "docker exec kobo-calibre-browser env | grep CALIBRE",
    "",
    "# Tester l'accès direct à la base SQLite depuis le container",
    'docker exec kobo-calibre-browser python -c "from database import test_connection; print(test_connection())"',
    "",
    "# Inspecter le réseau Docker",
    "docker network inspect docker_network | grep kobo",
    "",
    "# Forcer la reconstruction de l'image",
    "docker compose down && docker compose up -d --build --no-cache",
], "Commandes de diagnostic")

heading2(doc, "Mise à jour de l'application")
code_block(doc, [
    "cd /mnt/user/appdata/kobo-calibre-browser",
    "",
    "# Récupérer les dernières modifications",
    "git pull",
    "",
    "# Reconstruire et redémarrer",
    "docker compose down",
    "docker compose up -d --build",
], "Procédure de mise à jour")

# ── PIED DE PAGE ──────────────────────────────────────────────────────────────
section = doc.sections[0]
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.paragraph_format.space_before = Pt(4)
r_footer = fp.add_run("Kobo Calibre Browser — Guide d'installation v1.0  ·  fabien.majurel@outlook.com")
r_footer.font.size = Pt(8)
r_footer.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
r_footer.font.name = "Calibri"
r_footer.font.italic = True

# ── SAUVEGARDE ────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"OK  Document genere : {OUTPUT}")
